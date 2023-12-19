import pandas as pd
import docx
import numpy as np
import streamlit as st
from PIL import Image
from docx.shared import Inches, Pt, Mm
from docx.oxml import OxmlElement, ns
from docxtpl import DocxTemplate

def line_space(doc, number_of_lines):
    for i in range(number_of_lines):
        p = doc.add_paragraph()
        p.style = None
        p.add_run(style = None)
        p.add_run(style = None)

def heading_number_generaator(heading, sub_heading, sub_sub_heading):
    if sub_heading == 0:
        return str(heading)
    elif sub_sub_heading == 0:
        return str(heading)+'.'+str(sub_heading)
    else:
        return str(heading)+'.'+str(sub_heading)+'.'+str(sub_sub_heading)
    
def ahn(hn, level):
    headings = hn.split('.')
    if hn.count('.') == 2:
        heading = int(headings[0])
        sub_heading = int(headings[1])
        sub_sub_heading = int(headings[2])
    elif hn.count('.') == 1:
        heading = int(headings[0])
        sub_heading = int(headings[1])
        sub_sub_heading = 0
    else:
        heading = int(headings[0])
        sub_heading = 0
        sub_sub_heading = 0
    if level == 1:
        heading+=1
        return heading_number_generaator(heading,0,0)
    elif level == 2:
        sub_heading+=1
        return heading_number_generaator(heading,sub_heading,0)
    elif level == 3:
        sub_sub_heading+=1
        return heading_number_generaator(heading,sub_heading,sub_sub_heading)
    else:
        return(hn)

def add_site_observation_to_doc(string,doc,hn):
    rdict = {}
    n=0
    for i in string.split('\n'):
        if len(i)>0:
            if i[0] == '#':
                rdict[f'{n}_h'] = i[2:]
            elif i[0] == '$':
                rdict[f'{n}_s'] = i[2:]
            else:
                rdict[f'{n}_p'] = i
        n+=1
    for k,v in rdict.items():
        if len(v)>0:
            if k[-1] == 'h':
                hn = ahn(hn,2)
                doc.add_heading(f'{hn}. {v}',2)
                line_space(doc,1)
            elif k[-1] == 's':
                hn = ahn(hn,3)
                doc.add_heading(f'{hn}. {v}',3)
                line_space(doc,1)
            else:
                current_point = ''
                line = v.strip()
                if line.startswith('-'):
                    if current_point:
                        doc.add_paragraph(current_point.strip()[2:], style = 'List Bullet')
                        line_space(doc,1)
                        current_point =  ''
                current_point += ' '+line
                if current_point:
                    doc.add_paragraph(current_point.strip()[2:], style = 'List Bullet')
                    line_space(doc,1)
    return hn

def table_of_contents(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(ns.qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(ns.qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(ns.qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field"
    fldChar2.append(fldChar3)
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(ns.qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p
    return document

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_table_to_document(df, document, style = 'Table Grid'):
    if len(df) > 0:
        table = document.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        table.style = style
        for j, column in enumerate(df.columns):
            table.cell(0, j).text = column
        for i, row in df.iterrows():
            for j, value in enumerate(row):
                table.cell(i + 1, j).text = str(value)

def update_docx_fields(doc):
    try:
        for field in doc.inline_shapes:
            if field._inline.graphic.graphicData.xml.startswith('<w:drawing'):
                # Update specific fields as needed, this example updates all date fields
                if field._inline.graphic.graphicData.xml.startswith('<w:drawing><wp:docPr'):
                    field.text = "New Date"  # Change the content to a new value
        return doc
    except Exception as e:
        print(f"Error: {e}")
    

def add_points(x, doc):
    lines = x.split('\n')
    points = []
    current_point = ''
    for line in lines:
        line = line.strip()
        if line.startswith('-'):
            if current_point:
                points.append(current_point.strip())
                current_point =  ''
        current_point += ' '+line
    if current_point:
        points.append(current_point.strip())
    for i in points:
        doc.add_paragraph(i[2:],style = 'List Bullet')
        line_space(doc,1)

def make_inspection_document(client_name, client_location, unit_number, client_code, inspection_date, equipment_name, tag_number, inspection_type, edited_df, result_and_conclusion, site_observation, overall_summary, thickness_details, scanning_details, shellwise_inspection, tower_drawing, shell_plate_pics, detailed_report):
    # Create document
    doc = docx.Document()
    paragraph_format = doc.styles['Normal'].paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    # add_page_number(doc.sections[0].footer.paragraphs[0].add_run())

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = docx.shared.Pt(11)

    obj_styles = doc.styles
    obj_charstyle = obj_styles.add_style('CommentsStyle', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(16)
    obj_font.name = 'Arial'

    obj_styles = doc.styles
    obj_charstyle = obj_styles.add_style('SmallText', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(8)
    obj_font.name = 'Arial'
        # header
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    header = section.first_page_header 
    table = header.add_table(1, 3,width=Inches(8))
    hdr_cells = table.rows[0].cells
    p = hdr_cells[0].paragraphs[0] 
    format = p.paragraph_format
    format.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
    format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(r'E:\Works\GitHub\Inspection-Document-Generator\Files\Client\Logo\ArunTech.jpg', width=Inches(.7))

    p = hdr_cells[1].paragraphs[0] 
    format = p.paragraph_format
    format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'{client_name}\n{client_location} \n UNIT: {unit_number}')

    p = hdr_cells[2].paragraphs[0] 
    format = p.paragraph_format
    format.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
    format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(r'E:\Works\GitHub\Inspection-Document-Generator\Files\Client\Logo\HP.png', width=Inches(0.7))
    table.style = 'Table Grid'

    p = header.add_paragraph()
    p.style = None
    p.add_run(style = None)
    p.add_run(style = None)

    table_0 = doc.add_table(1,2)
    p = table_0.rows[0].cells[0].paragraphs[0]
    format = p.paragraph_format
    # format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    inspection_date = pd.to_datetime(inspection_date)
    date_str = inspection_date.strftime('%B %Y/%d')
    date_components = date_str.split(' ')
    date_components[0] = date_components[0][:3]
    date_components[0] = date_components[0].upper()
    datex = ' '.join(date_components)
    report_number = f'ATL/UT/{client_code}-{client_location}/{datex}'
    p.add_run('Report no:\n', style = 'SmallText')
    p.add_run(f"{report_number}")
    p = table_0.rows[0].cells[1].paragraphs[0]
    format = p.paragraph_format

    inspection_dateX = inspection_date.strftime('%d-%m-%Y')
    p.add_run('Inspection date:\n',style='SmallText')
    p.add_run(f'{inspection_dateX}')
    # format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    table_0.style = 'Colorful List'

    font.size = docx.shared.Pt(16)
    table_1 = doc.add_table(2,1)
    p = table_1.rows[0].cells[0].paragraphs[0]
    format = p.paragraph_format
    format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'{equipment_name}', style = 'CommentsStyle').bold = False
    p = table_1.rows[1].cells[0].paragraphs[0]
    format = p.paragraph_format
    format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'{inspection_type}', style = 'CommentsStyle').bold = True
    table_1.style = 'Colorful List'
    font.size = docx.shared.Pt(11)

    line_space(doc,30)
    edited_df.to_csv(rf'Temp\authors.csv')
    authors = pd.read_csv(rf'Temp\authors.csv')
    authors['Date'] = pd.to_datetime(authors['Date'])
    authors['Date'] = authors['Date'].dt.strftime('%d-%m-%Y')
    authors.drop('Unnamed: 0', inplace=True, axis=1)
    add_table_to_document(authors, doc, 'Colorful List')
    doc.add_page_break()

    section = doc.sections[0]
    header = section.header 
    table = header.add_table(2, 3,width=Inches(8))
    hdr_cells = table.rows[0].cells
    
    p = hdr_cells[0].paragraphs[0] 
    format = p.paragraph_format
    format.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
    format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(r'E:\Works\GitHub\Inspection-Document-Generator\Files\Client\Logo\ArunTech.jpg', width=Inches(.7))

    p = hdr_cells[1].paragraphs[0] 
    format = p.paragraph_format
    format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'{equipment_name}\nUNIT: {unit_number}\n')
    p.add_run(f'(Report No. {report_number})', style = 'SmallText')

    p = hdr_cells[2].paragraphs[0] 
    format = p.paragraph_format
    format.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
    format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(r'E:\Works\GitHub\Inspection-Document-Generator\Files\Client\Logo\HP.png', width=Inches(.7))

    sec_cells = table.rows[1].cells
    p = sec_cells[0].paragraphs[0] 
    format = p.paragraph_format
    format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Client:\n{client_code}-{unit_number}\n{client_location}")

    p = sec_cells[1].paragraphs[0] 
    format = p.paragraph_format
    format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{inspection_type}")

    p = sec_cells[2].paragraphs[0] 
    format = p.paragraph_format
    format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(p.add_run())

    table.style = 'Table Grid'

    p = header.add_paragraph()
    p.style = None
    p.add_run(style = None)
    p.add_run(style = None)

    p = doc.add_paragraph()
    p.add_run('Table of contents', style = 'CommentsStyle').bold = True
    format = p.paragraph_format
    format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    doc = table_of_contents(doc)
    doc.add_page_break()
    hn = '1'
    doc.add_heading(f'{hn}. Results and conclusion',1)
    line_space(doc,1)

    paragraph_format = doc.styles['List Bullet'].paragraph_format
    paragraph_format.left_indent = Inches(0.5)

    add_points(result_and_conclusion, doc)
    doc.add_page_break()

    hn = ahn(hn,1)
    doc.add_heading(f'{hn}. Site observation',1)
    line_space(doc,1)
    add_site_observation_to_doc(site_observation,doc,hn)
    doc.add_page_break()

    if bool(overall_summary)==True:
        hn = ahn(hn,1)
        doc.add_heading(f'{hn}. Overall Summary',1)
        line_space(doc,1)
        overall_summary_df = pd.read_csv(overall_summary,encoding='utf-8')
        add_table_to_document(overall_summary_df,doc)
        doc.add_page_break()

    if bool(thickness_details)==True:
        hn = ahn(hn,1)
        doc.add_heading(f'{hn}. Thickness Details',1)
        line_space(doc,1)
        overall_summary_df = pd.read_csv(thickness_details)
        add_table_to_document(overall_summary_df,doc)
        doc.add_page_break()

    if bool(scanning_details)==True:
        hn = ahn(hn,1)
        doc.add_heading(f'{hn}. Scanning Details',1)
        line_space(doc,1)
        overall_summary_df = pd.read_csv(scanning_details)
        add_table_to_document(overall_summary_df,doc)
        doc.add_page_break()

    if bool(shellwise_inspection)==True:
        hn = ahn(hn,1)
        doc.add_heading(f'{hn}. Shellwise Inspection',1)
        line_space(doc,1)
        for table in shellwise_inspection:
            overall_summary_df = pd.read_csv(table)
            add_table_to_document(overall_summary_df,doc)
            p = doc.add_paragraph()
            p.style = None
            p.add_run(style = None)
            p.add_run(style = None)
        doc.add_page_break()

    if bool(tower_drawing)==True:
        hn = ahn(hn,1)
        doc.add_heading(f'{hn}. Tower Drawings',1)
        line_space(doc,1)
        for pic in tower_drawing:
            image = Image.open(pic)
            image.save(rf'Temp\img.png')
            doc.add_picture(rf'Temp\img.png', width = Inches(5))
            p = doc.add_paragraph()
            p.style = None
            p.add_run(style = None)
            p.add_run(style = None)
        doc.add_page_break()

    if bool(shell_plate_pics)==True:
        hn = ahn(hn,1)
        doc.add_heading(f'{hn}. Shellplate Pictures',1)
        line_space(doc,1)
        for pic in shell_plate_pics:
            image = Image.open(pic)
            image.save(rf'Temp\img.png')
            doc.add_picture(rf'Temp\img.png', width = Inches(5))
            p = doc.add_paragraph()
            p.style = None
            p.add_run(style = None)
            p.add_run(style = None)
        doc.add_page_break()
    
    if bool(list(detailed_report.values())[0])==True:
        hn = ahn(hn,1)
        doc.add_heading(f'{hn}. Detailed Report',1)
        for k,v in detailed_report.items():
            hn = ahn(hn,2)
            doc.add_heading(f'{hn}. {k}',2)
            line_space(doc,1)
            for table in v:
                overall_summary_df = pd.read_csv(table)
                add_table_to_document(overall_summary_df,doc)
                p = doc.add_paragraph()
                p.style = None
                p.add_run(style = None)
                p.add_run(style = None)

    doc = update_docx_fields(doc)
    return doc
